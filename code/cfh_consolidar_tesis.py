# -*- coding: utf-8 -*-
r"""
cfh_consolidar_tesis.py
CFH - Hermeneutica Forense Computacional | Mireya Camacho Celis

Regenera el documento consolidado de la tesis UNIENDO las versiones FINAL de
los capitulos desde textos/. Garantiza consistencia total: el consolidado es
la union exacta de los capitulos ya revisados, sin editar contenido.

Estructura del consolidado:
  - Portada
  - (opcional) Indice
  - Cap 1..6 + Apendice + Referencias, con salto de pagina entre cada uno

Requiere: pip install docxcompose  (une docx preservando formato/estilos)
Si no esta instalado, el script lo instala automaticamente.

Uso (raiz del repo, env cfh):
    python code\cfh_consolidar_tesis.py

Salida: textos\CFH_Tesis_Consolidada_v8.docx
"""

import os
import sys
import subprocess

REPO = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
TEXTOS = os.path.join(REPO, "textos")
SALIDA = os.path.join(TEXTOS, "CFH_Tesis_Consolidada_v8.docx")

# Orden de union (nombres FINAL en textos/)
DOCUMENTOS = [
    "CFH_Tesis_Capitulo1_Introduccion_FINAL.docx",
    "CFH_Tesis_Capitulo2_EstadoDelArte_FINAL.docx",
    "CFH_Tesis_Capitulo3_MarcoTeorico_FINAL.docx",
    "CFH_Tesis_Capitulo4_Metodologia_FINAL.docx",
    "CFH_Tesis_Capitulo5_Resultados_FINAL.docx",
    "CFH_Tesis_Capitulo6_Discusion_FINAL.docx",
    "CFH_Apendice_Matematico_Experimentos_FINAL.docx",
    "CFH_Referencias_Consolidadas_APA7_FINAL.docx",
]

TITULO = ("El lenguaje de los falsos positivos: medición computacional multimodal "
          "de la injusticia discursiva y epistémica en el archivo judicial colombiano")
AUTORA = "Mireya Camacho Celis"
PROGRAMA = "Ciencia de Datos — Universidad Externado de Colombia"
DIRECTOR = "Director: Julián Zuluaga"
ANIO = "2026"


def asegurar_dependencias():
    try:
        import docxcompose  # noqa
        import docx  # noqa
        return True
    except ImportError:
        print("Instalando docxcompose y python-docx...")
        subprocess.check_call([sys.executable, "-m", "pip", "install",
                               "docxcompose", "python-docx", "--quiet"])
        return True


def crear_portada(path_portada):
    """Crea un docx con portada + indice como primer documento del merge."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
    from docx.enum.section import WD_SECTION

    doc = Document()

    # --- PORTADA ---
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(TITULO)
    r.bold = True; r.font.size = Pt(18)

    doc.add_paragraph()
    doc.add_paragraph()
    for texto, size, bold in [(AUTORA, 14, True), (PROGRAMA, 12, False),
                              (DIRECTOR, 12, False), (ANIO, 12, False)]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(texto); r.font.size = Pt(size); r.bold = bold

    doc.add_page_break()

    # --- INDICE (tabla de contenido simple, manual) ---
    p = doc.add_paragraph()
    r = p.add_run("Contenido"); r.bold = True; r.font.size = Pt(16)
    doc.add_paragraph()
    indice = [
        "1. Introducción",
        "2. Estado del arte",
        "3. Marco teórico",
        "4. Metodología",
        "5. Resultados",
        "6. Discusión y conclusiones",
        "Apéndice A — Desarrollo matemático y experimentos",
        "Referencias",
    ]
    for item in indice:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(6)

    doc.add_page_break()
    doc.save(path_portada)
    return path_portada


def main():
    print("=" * 64)
    print("CONSOLIDAR TESIS CFH - union de versiones FINAL")
    print("=" * 64)
    asegurar_dependencias()

    from docx import Document
    from docxcompose.composer import Composer

    # verificar que existan todos los FINAL
    faltan = []
    for d in DOCUMENTOS:
        if not os.path.exists(os.path.join(TEXTOS, d)):
            faltan.append(d)
    if faltan:
        print("\n[ERROR] Faltan documentos FINAL en textos/:")
        for f in faltan:
            print(f"    - {f}")
        print("\nGenera o copia los FINAL que faltan antes de consolidar.")
        return

    # crear portada
    portada = os.path.join(TEXTOS, "_portada_tmp.docx")
    crear_portada(portada)
    print(f"Portada + indice creados.")

    # componer: portada + cada capitulo (con salto de pagina entre ellos)
    master = Document(portada)
    composer = Composer(master)

    for i, nombre in enumerate(DOCUMENTOS):
        ruta = os.path.join(TEXTOS, nombre)
        sub = Document(ruta)
        # salto de pagina antes de cada documento (sobre el doc maestro directo)
        master.add_page_break()
        composer.append(sub)
        print(f"  [{i+1}/{len(DOCUMENTOS)}] unido: {nombre}")

    composer.save(SALIDA)

    # limpiar portada temporal
    try:
        os.remove(portada)
    except OSError:
        pass

    tam = os.path.getsize(SALIDA) / 1024
    print("\n" + "=" * 64)
    print(f"CONSOLIDADO GENERADO: {SALIDA}")
    print(f"Tamano: {tam:.0f} KB")
    print("=" * 64)
    print("\nEl consolidado v8 es la union EXACTA de los capitulos FINAL.")
    print("Consistente con opcion A, NV en IEI, surprisal descriptivo, IAA 0.722, B=80.")


if __name__ == "__main__":
    main()
